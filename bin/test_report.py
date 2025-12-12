import numpy as np
import os
import sys
import subprocess
import re
import random
from collections import defaultdict
import importlib.util
import time
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')

function_info = {
    'Quadratic': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function1.exe",
        'formula': 'x^2 + y^2 + z^2'
    },
    'Sinusoidal': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function2.exe",
        'formula': 'sin(x) + sin(y) + sin(z)'
    },
    'Exponential': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function3.exe",
        'formula': 'exp(0.1x) + exp(0.1y) + exp(0.1z)'
    },
    'Polynomial': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function4.exe",
        'formula': '1 + 0.5(x^2+y^2+z^2) + 0.1(x^3+y^3+z^3)'
    },
    'Rastrigin': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function5.exe",
        'formula': '30 + (x^2-10cos(6.28x)) + (y^2-10cos(6.28y)) + (z^2-10cos(6.28z))'
    },
    'Ackley': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function6.exe",
        'formula': '-20exp(-0.2sqrt((x^2+y^2+z^2)/3)) - exp((cos(6.28x)+cos(6.28y)+cos(6.28z))/3) + 20 + e'
    },
    'Product': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function7.exe",
        'formula': 'x * y * z'
    },
    'Trigonometric': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function8.exe",
        'formula': 'sin(x)cos(2x) + sin(y)cos(2y) + sin(z)cos(2z)'
    },
    'Logarithmic': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function9.exe",
        'formula': 'log(abs(x)+11) + log(abs(y)+11) + log(abs(z)+11)'
    },
    'Hyperbolic': {
        'input_dims': 3,
        'output_dims': 1,
        'exe_path': "function10.exe",
        'formula': 'tanh(x) + tanh(y) + tanh(z)'
    },
    'Linear_1D': {
        'input_dims': 1,
        'output_dims': 1,
        'exe_path': "function11.exe",
        'formula': '2x + 1'
    },
    'Quadratic_1D': {
        'input_dims': 1,
        'output_dims': 1,
        'exe_path': "function12.exe",
        'formula': 'x^2 + 2x + 1'
    },
    'Sin_1D': {
        'input_dims': 1,
        'output_dims': 1,
        'exe_path': "function13.exe",
        'formula': 'sin(6.28x)'
    },
    'Linear_2D': {
        'input_dims': 2,
        'output_dims': 1,
        'exe_path': "function14.exe",
        'formula': 'x + 2y + 1'
    },
    'Mixed_2D': {
        'input_dims': 2,
        'output_dims': 1,
        'exe_path': "function15.exe",
        'formula': 'x*y + sin(x) + cos(y)'
    },
    'Gaussian_2D': {
        'input_dims': 2,
        'output_dims': 1,
        'exe_path': "function16.exe",
        'formula': 'exp(-(x^2 + y^2)/2)'
    },
    'Vector_3D_2out': {
        'input_dims': 3,
        'output_dims': 2,
        'exe_path': "function17.exe",
        'formula': '[x + y, y * z]'
    },
    'Mixed_3D_2out': {
        'input_dims': 3,
        'output_dims': 2,
        'exe_path': "function18.exe",
        'formula': '[sin(x) + cos(y), z^2 + x*y]'
    },
    'HyperSphere_4D': {
        'input_dims': 4,
        'output_dims': 1,
        'exe_path': "function19.exe",
        'formula': 'x1^2 + x2^2 + x3^2 + x4^2'
    },
    'Linear_4D': {
        'input_dims': 4,
        'output_dims': 1,
        'exe_path': "function20.exe",
        'formula': 'x1 + 2x2 + 3x3 + 4x4 + 1'
    },
    'Vector_4D_2out': {
        'input_dims': 4,
        'output_dims': 2,
        'exe_path': "function21.exe",
        'formula': '[x1 + x2 + x3, x2 * x3 * x4]'
    },
    'Mixed_4D_2out': {
        'input_dims': 4,
        'output_dims': 2,
        'exe_path': "function22.exe",
        'formula': '[sin(x1) + cos(x2), exp(x3) + log(abs(x4)+1)]'
    },
    'Vector_4D_3out': {
        'input_dims': 4,
        'output_dims': 3,
        'exe_path': "function23.exe",
        'formula': '[x1, x2, x3 + x4]'
    },
    'Mixed_4D_3out': {
        'input_dims': 4,
        'output_dims': 3,
        'exe_path': "function24.exe",
        'formula': '[sin(x1), cos(x2), tanh(x3 + x4)]'
    },
    'HyperSphere_5D': {
        'input_dims': 5,
        'output_dims': 1,
        'exe_path': "function25.exe",
        'formula': 'x1^2 + x2^2 + x3^2 + x4^2 + x5^2'
    },
    'Linear_5D': {
        'input_dims': 5,
        'output_dims': 1,
        'exe_path': "function26.exe",
        'formula': 'x1 + 2x2 + 3x3 + 4x4 + 5x5 + 1'
    },
    'Vector_5D_2out': {
        'input_dims': 5,
        'output_dims': 2,
        'exe_path': "function27.exe",
        'formula': '[x1 + x2 + x3, x4 * x5 + x3]'
    },
    'Mixed_5D_2out': {
        'input_dims': 5,
        'output_dims': 2,
        'exe_path': "function28.exe",
        'formula': '[sin(x1) + cos(x2), exp(x3) + x4*x5]'
    },
    'Vector_5D_3out': {
        'input_dims': 5,
        'output_dims': 3,
        'exe_path': "function29.exe",
        'formula': '[x1 + x2, x3 * x4, x5^2]'
    },
    'Mixed_5D_3out': {
        'input_dims': 5,
        'output_dims': 3,
        'exe_path': "function30.exe",
        'formula': '[sin(x1+x2), cos(x3*x4), tanh(x5)]'
    },
}

models_info = [
    'Quadratic_rbf_50samples.py',
    'Quadratic_prbf_50samples.py',
    'Quadratic_rf_50samples.py',
    'Sinusoidal_rbf_50samples.py',
    'Sinusoidal_prbf_50samples.py',
    'Sinusoidal_rf_50samples.py',
    'Quadratic_xgboost_50samples.py',
    'Quadratic_prbf_100samples.py',
    'Quadratic_rbf_100samples.py',
    'Sinusoidal_rbf_100samples.py',
    'Quadratic_rf_100samples.py',
    'Sinusoidal_prbf_100samples.py',
    'Exponential_rbf_50samples.py',
    'Sinusoidal_rf_100samples.py',
    'Exponential_prbf_50samples.py',
    'Sinusoidal_xgboost_50samples.py',
    'Quadratic_nn_50samples.py',
    'Exponential_prbf_100samples.py',
    'Exponential_rbf_100samples.py',
    'Quadratic_rbf_200samples.py',
    'Exponential_rf_50samples.py',
    'Quadratic_prbf_200samples.py',
    'Sinusoidal_prbf_200samples.py',
    'Sinusoidal_rbf_200samples.py',
    'Sinusoidal_nn_50samples.py',
    'Exponential_xgboost_50samples.py',
    'Exponential_rf_100samples.py',
    'Polynomial_rbf_50samples.py',
    'Exponential_prbf_200samples.py',
    'Exponential_rbf_200samples.py',
    'Sinusoidal_rf_200samples.py',
    'Polynomial_rf_50samples.py',
    'Polynomial_prbf_50samples.py',
    'Polynomial_xgboost_50samples.py',
    'Polynomial_rbf_100samples.py',
    'Polynomial_prbf_100samples.py',
    'Quadratic_rf_200samples.py',
    'Exponential_nn_50samples.py',
    'Quadratic_xgboost_100samples.py',
    'Polynomial_rf_100samples.py',
    'Rastrigin_rbf_50samples.py',
    'Sinusoidal_nn_100samples.py',
    'Sinusoidal_xgboost_100samples.py',
    'Exponential_rf_200samples.py',
    'Polynomial_rbf_200samples.py',
    'Polynomial_prbf_200samples.py',
    'Polynomial_nn_50samples.py',
    'Rastrigin_rbf_100samples.py',
    'Rastrigin_prbf_50samples.py',
    'Rastrigin_rf_50samples.py',
    'Exponential_xgboost_100samples.py',
    'Rastrigin_prbf_100samples.py',
    'Rastrigin_xgboost_50samples.py',
    'Ackley_prbf_50samples.py',
    'Rastrigin_rf_100samples.py',
    'Ackley_rbf_50samples.py',
    'Rastrigin_prbf_200samples.py',
    'Rastrigin_rbf_200samples.py',
    'Ackley_rf_50samples.py',
    'Ackley_rf_100samples.py',
    'Rastrigin_nn_50samples.py',
    'Ackley_prbf_100samples.py',
    'Ackley_rbf_100samples.py',
    'Polynomial_rf_200samples.py',
    'Ackley_prbf_200samples.py',
    'Ackley_rbf_200samples.py',
    'Ackley_nn_50samples.py',
    'Ackley_xgboost_50samples.py',
    'Polynomial_xgboost_100samples.py',
    'Product_prbf_50samples.py',
    'Product_rbf_50samples.py',
    'Product_rbf_100samples.py',
    'Product_prbf_100samples.py',
    'Rastrigin_rf_200samples.py',
    'Ackley_rf_200samples.py',
    'Product_rbf_200samples.py',
    'Product_rf_50samples.py',
    'Quadratic_nn_100samples.py',
    'Product_xgboost_50samples.py',
    'Product_rf_100samples.py',
    'Rastrigin_xgboost_100samples.py',
    'Product_prbf_200samples.py',
    'Product_nn_50samples.py',
    'Trigonometric_rbf_50samples.py',
    'Trigonometric_rbf_100samples.py',
    'Trigonometric_prbf_50samples.py',
    'Exponential_nn_100samples.py',
    'Ackley_xgboost_100samples.py',
    'Trigonometric_rbf_200samples.py',
    'Trigonometric_prbf_100samples.py',
    'Trigonometric_rf_50samples.py',
    'Polynomial_nn_100samples.py',
    'Logarithmic_rbf_50samples.py',
    'Trigonometric_rf_100samples.py',
    'Trigonometric_xgboost_50samples.py',
    'Trigonometric_prbf_200samples.py',
    'Product_rf_200samples.py',
    'Trigonometric_nn_50samples.py',
    'Logarithmic_rbf_100samples.py',
    'Rastrigin_nn_100samples.py',
    'Product_xgboost_100samples.py',
    'Logarithmic_prbf_50samples.py',
    'Trigonometric_rf_200samples.py',
    'Logarithmic_prbf_100samples.py',
    'Logarithmic_rbf_200samples.py',
    'Logarithmic_rf_50samples.py',
    'Trigonometric_xgboost_100samples.py',
    'Logarithmic_xgboost_50samples.py',
    'Logarithmic_rf_100samples.py',
    'Trigonometric_nn_100samples.py',
    'Logarithmic_prbf_200samples.py',
    'Ackley_nn_100samples.py',
    'Logarithmic_nn_50samples.py',
    'Hyperbolic_rbf_100samples.py',
    'Hyperbolic_rbf_50samples.py',
    'Hyperbolic_rf_50samples.py',
    'Hyperbolic_prbf_50samples.py',
    'Logarithmic_rf_200samples.py',
    'Hyperbolic_prbf_100samples.py',
    'Product_nn_100samples.py',
    'Hyperbolic_rf_100samples.py',
    'Hyperbolic_nn_50samples.py',
    'Hyperbolic_rbf_200samples.py',
    'Hyperbolic_prbf_200samples.py',
    'Logarithmic_xgboost_100samples.py',
    'Hyperbolic_xgboost_50samples.py',
    'Hyperbolic_rf_200samples.py',
    'Linear_1D_rbf_50samples.py',
    'Linear_1D_prbf_50samples.py',
    'Linear_1D_rf_50samples.py',
    'Linear_1D_rbf_100samples.py',
    'Linear_1D_prbf_100samples.py',
    'Linear_1D_rf_100samples.py',
    'Linear_1D_rbf_200samples.py',
    'Linear_1D_prbf_200samples.py',
    'Hyperbolic_xgboost_100samples.py',
    'Linear_1D_nn_50samples.py',
    'Linear_1D_rf_200samples.py',
    'Logarithmic_nn_100samples.py',
    'Sinusoidal_xgboost_200samples.py',
    'Hyperbolic_nn_100samples.py',
    'Linear_1D_xgboost_50samples.py',
    'Exponential_xgboost_200samples.py',
    'Quadratic_1D_rbf_50samples.py',
    'Quadratic_1D_rbf_100samples.py',
    'Linear_1D_nn_100samples.py',
    'Quadratic_1D_prbf_50samples.py',
    'Quadratic_1D_rbf_200samples.py',
    'Quadratic_1D_rf_50samples.py',
    'Quadratic_1D_prbf_100samples.py',
    'Quadratic_1D_rf_100samples.py',
    'Linear_1D_xgboost_100samples.py',
    'Quadratic_1D_prbf_200samples.py',
    'Quadratic_1D_nn_50samples.py',
    'Quadratic_1D_rf_200samples.py',
    'Quadratic_1D_xgboost_50samples.py',
    'Quadratic_xgboost_200samples.py',
    'Sin_1D_rbf_50samples.py',
    'Sin_1D_rbf_100samples.py',
    'Quadratic_1D_nn_100samples.py',
    'Quadratic_1D_xgboost_100samples.py',
    'Sin_1D_prbf_50samples.py',
    'Sin_1D_rbf_200samples.py',
    'Sin_1D_prbf_100samples.py',
    'Sin_1D_rf_50samples.py',
    'Sin_1D_prbf_200samples.py',
    'Sin_1D_rf_100samples.py',
    'Ackley_xgboost_200samples.py',
    'Polynomial_xgboost_200samples.py',
    'Sin_1D_nn_50samples.py',
    'Sin_1D_xgboost_50samples.py',
    'Sin_1D_nn_100samples.py',
    'Sin_1D_rf_200samples.py',
    'Linear_2D_rbf_50samples.py',
    'Rastrigin_xgboost_200samples.py',
    'Linear_2D_rbf_100samples.py',
    'Linear_1D_nn_200samples.py',
    'Trigonometric_xgboost_200samples.py',
    'Linear_2D_prbf_50samples.py',
    'Sin_1D_xgboost_100samples.py',
    'Linear_1D_xgboost_200samples.py',
    'Linear_2D_prbf_100samples.py',
    'Linear_2D_rbf_200samples.py',
    'Linear_2D_rf_50samples.py',
    'Logarithmic_xgboost_200samples.py',
    'Linear_2D_rf_100samples.py',
    'Mixed_2D_rbf_50samples.py',
    'Mixed_2D_prbf_50samples.py',
    'Mixed_2D_rf_50samples.py',
    'Linear_2D_prbf_200samples.py',
    'Linear_2D_xgboost_50samples.py',
    'Mixed_2D_rf_100samples.py',
    'Linear_2D_nn_50samples.py',
    'Mixed_2D_rbf_100samples.py',
    'Mixed_2D_prbf_100samples.py',
    'Mixed_2D_prbf_200samples.py',
    'Mixed_2D_rbf_200samples.py',
    'Linear_2D_rf_200samples.py',
    'Mixed_2D_nn_50samples.py',
    'Gaussian_2D_rbf_50samples.py',
    'Quadratic_1D_nn_200samples.py',
    'Mixed_2D_xgboost_50samples.py',
    'Product_xgboost_200samples.py',
    'Linear_2D_xgboost_100samples.py',
    'Quadratic_1D_xgboost_200samples.py',
    'Gaussian_2D_rf_50samples.py',
    'Gaussian_2D_prbf_50samples.py',
    'Sin_1D_nn_200samples.py',
    'Hyperbolic_xgboost_200samples.py',
    'Linear_2D_nn_100samples.py',
    'Gaussian_2D_rf_100samples.py',
    'Mixed_2D_nn_100samples.py',
    'Mixed_2D_rf_200samples.py',
    'Gaussian_2D_prbf_100samples.py',
    'Gaussian_2D_rbf_100samples.py',
    'Gaussian_2D_xgboost_50samples.py',
    'Gaussian_2D_nn_50samples.py',
    'Vector_3D_2out_rbf_50samples.py',
    'Gaussian_2D_prbf_200samples.py',
    'Gaussian_2D_rbf_200samples.py',
    'Vector_3D_2out_rbf_100samples.py',
    'Sin_1D_xgboost_200samples.py',
    'Mixed_2D_xgboost_100samples.py',
    'Vector_3D_2out_rf_50samples.py',
    'Vector_3D_2out_prbf_50samples.py',
    'Vector_3D_2out_prbf_100samples.py',
    'Vector_3D_2out_rf_100samples.py',
    'Vector_3D_2out_rbf_200samples.py',
    'Gaussian_2D_xgboost_100samples.py',
    'Vector_3D_2out_prbf_200samples.py',
    'Gaussian_2D_rf_200samples.py',
    'Mixed_3D_2out_rbf_50samples.py',
    'Mixed_3D_2out_rbf_100samples.py',
    'Vector_3D_2out_nn_50samples.py',
    'Vector_3D_2out_xgboost_50samples.py',
    'Gaussian_2D_nn_100samples.py',
    'Mixed_3D_2out_prbf_50samples.py',
    'Mixed_3D_2out_rbf_200samples.py',
    'Mixed_3D_2out_rf_50samples.py',
    'Mixed_3D_2out_prbf_100samples.py',
    'Mixed_3D_2out_prbf_200samples.py',
    'Vector_3D_2out_rf_200samples.py',
    'Mixed_3D_2out_xgboost_50samples.py',
    'Mixed_3D_2out_rf_100samples.py',
    'Mixed_3D_2out_nn_50samples.py',
    'Sinusoidal_nn_200samples.py',
    'Mixed_3D_2out_rf_200samples.py',
    'HyperSphere_4D_rbf_50samples.py',
    'HyperSphere_4D_rbf_100samples.py',
    'HyperSphere_4D_prbf_50samples.py',
    'Mixed_2D_nn_200samples.py',
    'Vector_3D_2out_xgboost_100samples.py',
    'HyperSphere_4D_rbf_200samples.py',
    'Gaussian_2D_nn_200samples.py',
    'HyperSphere_4D_rf_50samples.py',
    'HyperSphere_4D_prbf_100samples.py',
    'HyperSphere_4D_rf_100samples.py',
    'HyperSphere_4D_prbf_200samples.py',
    'Linear_2D_xgboost_200samples.py',
    'HyperSphere_4D_nn_50samples.py',
    'Linear_2D_nn_200samples.py',
    'Gaussian_2D_xgboost_200samples.py',
    'Linear_4D_rbf_50samples.py',
    'Vector_3D_2out_nn_100samples.py',
    'Quadratic_nn_200samples.py',
    'HyperSphere_4D_xgboost_50samples.py',
    'Linear_4D_prbf_50samples.py',
    'Mixed_3D_2out_xgboost_100samples.py',
    'Linear_4D_rbf_100samples.py',
    'Linear_4D_prbf_100samples.py',
    'Linear_4D_rf_50samples.py',
    'Linear_4D_xgboost_50samples.py',
    'HyperSphere_4D_rf_200samples.py',
    'Linear_4D_rf_100samples.py',
    'Linear_4D_rbf_200samples.py',
    'Linear_4D_prbf_200samples.py',
    'Mixed_3D_2out_nn_100samples.py',
    'Exponential_nn_200samples.py',
    'Vector_4D_2out_rbf_50samples.py',
    'Vector_4D_2out_rbf_100samples.py',
    'Vector_4D_2out_prbf_50samples.py',
    'Vector_4D_2out_prbf_100samples.py',
    'Polynomial_nn_200samples.py',
    'Linear_4D_nn_50samples.py',
    'Mixed_2D_xgboost_200samples.py',
    'Linear_4D_rf_200samples.py',
    'Vector_4D_2out_rbf_200samples.py',
    'Vector_4D_2out_rf_50samples.py',
    'Trigonometric_nn_200samples.py',
    'Vector_4D_2out_rf_100samples.py',
    'HyperSphere_4D_xgboost_100samples.py',
    'Vector_4D_2out_prbf_200samples.py',
    'Rastrigin_nn_200samples.py',
    'Ackley_nn_200samples.py',
    'Vector_4D_2out_nn_50samples.py',
    'Vector_4D_2out_xgboost_50samples.py',
    'Mixed_4D_2out_rbf_50samples.py',
    'Mixed_4D_2out_prbf_50samples.py',
    'Mixed_4D_2out_rbf_100samples.py',
    'Mixed_4D_2out_prbf_100samples.py',
    'Mixed_4D_2out_rf_50samples.py',
    'Product_nn_200samples.py',
    'Linear_4D_xgboost_100samples.py',
    'Mixed_4D_2out_rbf_200samples.py',
    'Mixed_4D_2out_nn_50samples.py',
    'Mixed_4D_2out_prbf_200samples.py',
    'Mixed_4D_2out_rf_100samples.py',
    'Vector_4D_3out_rbf_50samples.py',
    'Vector_4D_3out_rbf_100samples.py',
    'Mixed_4D_2out_xgboost_50samples.py',
    'Vector_4D_3out_prbf_50samples.py',
    'Vector_4D_3out_prbf_100samples.py',
    'HyperSphere_4D_nn_100samples.py',
    'Vector_4D_3out_rf_50samples.py',
    'Vector_4D_3out_rbf_200samples.py',
    'Vector_4D_2out_rf_200samples.py',
    'Vector_4D_3out_prbf_200samples.py',
    'Vector_4D_3out_rf_100samples.py',
    'Hyperbolic_nn_200samples.py',
    'Mixed_4D_2out_rf_200samples.py',
    'Vector_4D_2out_nn_100samples.py',
    'Mixed_4D_3out_rbf_50samples.py',
    'Linear_4D_nn_100samples.py',
    'Vector_4D_3out_xgboost_50samples.py',
    'Vector_4D_3out_nn_50samples.py',
    'Mixed_4D_3out_prbf_50samples.py',
    'Mixed_4D_3out_rf_50samples.py',
    'Vector_4D_2out_xgboost_100samples.py',
    'Mixed_4D_3out_rbf_100samples.py',
    'Mixed_4D_3out_prbf_100samples.py',
    'Logarithmic_nn_200samples.py',
    'Mixed_4D_3out_nn_50samples.py',
    'Mixed_4D_3out_xgboost_50samples.py',
    'Mixed_4D_3out_rf_100samples.py',
    'Mixed_4D_3out_rbf_200samples.py',
    'Mixed_4D_3out_prbf_200samples.py',
    'HyperSphere_5D_rbf_50samples.py',
    'Mixed_4D_2out_xgboost_100samples.py',
    'Vector_4D_3out_rf_200samples.py',
    'HyperSphere_5D_rf_50samples.py',
    'HyperSphere_5D_prbf_50samples.py',
    'HyperSphere_5D_rbf_100samples.py',
    'HyperSphere_5D_prbf_100samples.py',
    'HyperSphere_5D_prbf_200samples.py',
    'HyperSphere_5D_rbf_200samples.py',
    'Mixed_4D_2out_nn_100samples.py',
    'HyperSphere_5D_rf_100samples.py',
    'Mixed_4D_3out_rf_200samples.py',
    'HyperSphere_5D_nn_50samples.py',
    'HyperSphere_5D_rf_200samples.py',
    'HyperSphere_5D_xgboost_50samples.py',
    'Mixed_4D_3out_nn_100samples.py',
    'Linear_5D_rbf_50samples.py',
    'Linear_5D_rbf_100samples.py',
    'Linear_5D_prbf_50samples.py',
    'Linear_5D_rf_50samples.py',
    'Vector_4D_3out_xgboost_100samples.py',
    'Linear_5D_prbf_100samples.py',
    'Vector_4D_3out_nn_100samples.py',
    'Linear_5D_rf_100samples.py',
    'Linear_5D_rbf_200samples.py',
    'Linear_5D_prbf_200samples.py',
    'Mixed_4D_3out_xgboost_100samples.py',
    'Linear_5D_rf_200samples.py',
    'Linear_5D_nn_50samples.py',
    'Vector_5D_2out_prbf_50samples.py',
    'Vector_5D_2out_rbf_50samples.py',
    'Vector_5D_2out_rbf_100samples.py',
    'Linear_5D_xgboost_50samples.py',
    'HyperSphere_5D_xgboost_100samples.py',
    'Vector_5D_2out_rbf_200samples.py',
    'Vector_5D_2out_prbf_100samples.py',
    'Vector_5D_2out_rf_50samples.py',
    'Vector_5D_2out_rf_100samples.py',
    'Vector_5D_2out_prbf_200samples.py',
    'Vector_5D_2out_xgboost_50samples.py',
    'Linear_5D_xgboost_100samples.py',
    'Vector_5D_2out_nn_50samples.py',
    'HyperSphere_5D_nn_100samples.py',
    'Linear_4D_xgboost_200samples.py',
    'Vector_5D_2out_rf_200samples.py',
    'HyperSphere_4D_xgboost_200samples.py',
    'Mixed_5D_2out_prbf_50samples.py',
    'Mixed_5D_2out_rbf_50samples.py',
    'Mixed_5D_2out_rbf_100samples.py',
    'Mixed_5D_2out_rbf_200samples.py',
    'Vector_3D_2out_xgboost_200samples.py',
    'Mixed_5D_2out_rf_50samples.py',
    'Mixed_5D_2out_prbf_100samples.py',
    'Linear_5D_nn_100samples.py',
    'Mixed_5D_2out_nn_50samples.py',
    'Mixed_3D_2out_xgboost_200samples.py',
    'Mixed_5D_2out_rf_100samples.py',
    'Mixed_5D_2out_prbf_200samples.py',
    'Vector_5D_3out_rbf_50samples.py',
    'Vector_5D_3out_prbf_50samples.py',
    'Vector_5D_2out_xgboost_100samples.py',
    'Vector_5D_3out_rbf_100samples.py',
    'Vector_5D_3out_prbf_100samples.py',
    'Mixed_5D_2out_xgboost_50samples.py',
    'Vector_5D_3out_rbf_200samples.py',
    'Mixed_5D_2out_rf_200samples.py',
    'Mixed_5D_2out_nn_100samples.py',
    'Vector_5D_3out_rf_50samples.py',
    'Vector_5D_3out_xgboost_50samples.py',
    'Vector_5D_3out_rf_100samples.py',
    'Vector_5D_2out_nn_100samples.py',
    'Vector_5D_3out_prbf_200samples.py',
    'Vector_3D_2out_nn_200samples.py',
    'Vector_5D_3out_nn_50samples.py',
    'Mixed_3D_2out_nn_200samples.py',
    'Mixed_5D_2out_xgboost_100samples.py',
    'Mixed_5D_3out_rbf_50samples.py',
    'Mixed_5D_3out_prbf_50samples.py',
    'Mixed_5D_3out_rbf_100samples.py',
    'Vector_5D_3out_rf_200samples.py',
    'Mixed_5D_3out_rbf_200samples.py',
    'Mixed_5D_3out_prbf_100samples.py',
    'Mixed_5D_3out_rf_50samples.py',
    'Mixed_5D_3out_prbf_200samples.py',
    'Mixed_5D_3out_rf_100samples.py',
    'Mixed_5D_3out_xgboost_50samples.py',
    'Mixed_5D_3out_nn_50samples.py',
    'Vector_5D_3out_xgboost_100samples.py',
    'Mixed_5D_3out_rf_200samples.py',
    'Vector_5D_3out_nn_100samples.py',
    'HyperSphere_5D_xgboost_200samples.py',
    'Linear_5D_xgboost_200samples.py',
    'Vector_4D_2out_xgboost_200samples.py',
    'HyperSphere_4D_nn_200samples.py',
    'Mixed_5D_3out_nn_100samples.py',
    'Linear_4D_nn_200samples.py',
    'Mixed_4D_2out_xgboost_200samples.py',
    'Mixed_5D_3out_xgboost_100samples.py',
    'Vector_4D_2out_nn_200samples.py',
    'Mixed_4D_2out_nn_200samples.py',
    'Mixed_4D_3out_xgboost_200samples.py',
    'Vector_4D_3out_xgboost_200samples.py',
    'Mixed_4D_3out_nn_200samples.py',
    'Vector_4D_3out_nn_200samples.py',
    'Vector_5D_2out_xgboost_200samples.py',
    'HyperSphere_5D_nn_200samples.py',
    'Mixed_5D_2out_xgboost_200samples.py',
    'Linear_5D_nn_200samples.py',
    'Mixed_5D_2out_nn_200samples.py',
    'Vector_5D_2out_nn_200samples.py',
    'Mixed_5D_3out_xgboost_200samples.py',
    'Mixed_5D_3out_nn_200samples.py',
    'Vector_5D_3out_xgboost_200samples.py',
    'Vector_5D_3out_nn_200samples.py',
]

_model_cache = {}

def get_true_values(exe_path, input_point):
    """Call the executable with input point and get true output values"""
    try:
        # Build command
        cmd = [exe_path] + [str(x) for x in input_point]
        
        # Run process
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
        
        if result.returncode != 0:
            print(f"  ERROR: Process failed with return code {result.returncode}")
            return None
        
        # Parse output - can be single value or multiple lines for multiple outputs
        output_lines = result.stdout.strip().split('\n')
        output_values = []
        
        for line in output_lines:
            if line.strip():
                try:
                    value = float(line.strip())
                    output_values.append(value)
                except ValueError:
                    print(f"  WARNING: Could not parse output line: '{line}'")
        
        return output_values if output_values else None
        
    except subprocess.TimeoutExpired:
        print(f"  ERROR: Process timeout for {exe_path}")
        return None
    except Exception as e:
        print(f"  ERROR: {e}")
        return None

def load_model(model_file):
    """Load model from file with caching"""
    if model_file in _model_cache:
        return _model_cache[model_file]
    
    try:
        model_dir = os.path.dirname(model_file)
        model_name = os.path.basename(model_file).replace('.py', '')
        
        # Add model directory to Python path
        if model_dir and model_dir not in sys.path:
            sys.path.insert(0, model_dir)
        
        # Import the model module
        spec = importlib.util.spec_from_file_location(model_name, model_file)
        if spec is None:
            print(f"  ERROR: Could not load spec from {model_file}")
            return None
            
        model_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(model_module)
        
        model_instance = None
        
        if hasattr(model_module, 'PRBFModel'):
            model_instance = model_module.PRBFModel()
        elif hasattr(model_module, 'RBFModel'):
            try:
                model_instance = model_module.RBFModel()
            except TypeError:
                try:
                    model_instance = model_module.RBFModel([], 1.0)
                except:
                    print(f"  ERROR: Could not initialize RBFModel from {model_file}")
                    return None
        elif hasattr(model_module, 'NNModel'):
            model_instance = model_module.NNModel()
        elif hasattr(model_module, 'RFModel'):
            model_instance = model_module.RFModel()
        elif hasattr(model_module, 'XGBoostModel'):
            model_instance = model_module.XGBoostModel()
        
        if model_instance is None:
            print(f"  ERROR: No compatible model class found in {model_file}")
            return None
        
        _model_cache[model_file] = model_instance
        return model_instance
        
    except Exception as e:
        print(f"  ERROR loading model {model_file}: {e}")
        import traceback
        traceback.print_exc()
        return None
        
    except Exception as e:
        print(f"  ERROR loading model {model_file}: {e}")
        return None

def evaluate_model_prediction(model_file, input_point, true_values):
    """Evaluate model prediction using cached model"""
    try:
        model_instance = load_model(model_file)
        if model_instance is None:
            return None
        
        # Get prediction
        predicted_values = model_instance.predict(input_point)
        
        # Convert to list if single value
        if predicted_values is not None:
            if not hasattr(predicted_values, '__len__') or isinstance(predicted_values, (int, float)):
                predicted_values = [predicted_values]
            elif isinstance(predicted_values, np.ndarray):
                predicted_values = predicted_values.tolist()
            elif isinstance(predicted_values, tuple):
                predicted_values = list(predicted_values)
        
        return predicted_values if predicted_values is not None else None
        
    except Exception as e:
        print(f"  ERROR evaluating model {model_file}: {e}")
        return None

def calculate_error(true_values, predicted_values):
    """Calculate mean absolute error between true and predicted values"""
    if true_values is None or predicted_values is None:
        return float('inf')
    
    if len(true_values) != len(predicted_values):
        print(f"  WARNING: Dimension mismatch: true={len(true_values)}, pred={len(predicted_values)}")
        return float('inf')
    
    # Calculate mean absolute error across all outputs
    errors = [abs(t - p) for t, p in zip(true_values, predicted_values)]
    mean_error = sum(errors) / len(errors)
    
    return mean_error

def parse_model_filename(model_file):
    """Extract function name, model type, and sample size from filename"""
    base_name = os.path.basename(model_file).replace('.py', '')
    parts = base_name.split('_')
    
    if len(parts) >= 3:
        # Reconstruct function name from all parts except last two
        func_name_parts = parts[:-2]
        func_name = '_'.join(func_name_parts)
        model_type = parts[-2]
        samples_str = parts[-1].replace('samples', '')
    else:
        func_name = 'Unknown'
        model_type = 'unknown'
        samples_str = '0'
    
    try:
        samples = int(samples_str)
    except:
        samples = 0
    
    return func_name, model_type, samples

def test_single_model(model_file, num_test_points=1000):
    """Test a single model on random points and calculate average error"""
    func_name, model_type, samples = parse_model_filename(model_file)
    
    if func_name not in function_info:
        print(f"  WARNING: Unknown function {func_name}")
        return float('inf')
    
    func_data = function_info[func_name]
    input_dims = func_data['input_dims']
    exe_path = func_data['exe_path']
    
    if not os.path.exists(exe_path):
        print(f"  ERROR: Executable not found: {exe_path}")
        return float('inf')
    
    print(f"  Testing {func_name} - {model_type} ({samples} samples)...")
    
    total_error = 0.0
    successful_tests = 0
    
    model_instance = load_model(model_file)
    if model_instance is None:
        return float('inf')
    
    for i in range(num_test_points):
        # Generate random input point in range [-10, 10]
        input_point = [random.uniform(-10.0, 10.0) for _ in range(input_dims)]
        
        # Get true values from executable
        true_values = get_true_values(exe_path, input_point)
        if true_values is None:
            continue
        
        # Get predicted values from model
        try:
            predicted_values = model_instance.predict(input_point)
            
            # Convert to list if single value
            if predicted_values is not None:
                if not hasattr(predicted_values, '__len__') or isinstance(predicted_values, (int, float)):
                    predicted_values = [predicted_values]
                elif isinstance(predicted_values, np.ndarray):
                    predicted_values = predicted_values.tolist()
                elif isinstance(predicted_values, tuple):
                    predicted_values = list(predicted_values)
        except Exception as e:
            print(f"    Prediction error: {e}")
            continue
        
        if predicted_values is None:
            continue
        
        # Calculate error
        error = calculate_error(true_values, predicted_values)
        if error < float('inf'):
            total_error += error
            successful_tests += 1
    
    if successful_tests == 0:
        print(f"    No successful tests for {model_file}")
        return float('inf')
    
    avg_error = total_error / successful_tests
    print(f"    Average error: {avg_error:.6f} ({successful_tests}/{num_test_points} successful tests)")
    
    return avg_error

def create_word_report(results, total_time):
    """Create Word document with tables and graphs"""
    doc = Document()
    
    # Title
    doc.add_heading('AI Models Test Report - Error Analysis', 0)
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f'Total testing time: {total_time:.2f} seconds')
    
    # Calculate overall statistics
    total_models = len(models_info)
    successful_models = 0
    total_error = 0.0
    
    for func_name, model_data in results.items():
        for model_type, samples_data in model_data.items():
            for samples, error in samples_data.items():
                if error < float('inf'):
                    successful_models += 1
                    total_error += error
    
    doc.add_paragraph(f'Total models tested: {total_models}')
    doc.add_paragraph(f'Successful models: {successful_models}')
    
    if successful_models > 0:
        avg_error = total_error / successful_models
        doc.add_paragraph(f'Average error across all models: {avg_error:.6f}')
    
    # Results for each function
    doc.add_heading('Detailed Results by Function', level=1)
    
    for func_name, model_data in sorted(results.items()):
        if func_name not in function_info:
            continue
            
        func_info = function_info[func_name]
        input_dims = func_info['input_dims']
        output_dims = func_info['output_dims']
        formula = func_info['formula']
        
        # Function header
        doc.add_heading(f'Function: {func_name}', level=2)
        doc.add_paragraph(f'Dimensions: {input_dims}D -> {output_dims}D')
        doc.add_paragraph(f'Formula: {formula}')
        
        # Create table
        sample_sizes = set()
        model_types = set()
        
        for model_type, samples_data in model_data.items():
            model_types.add(model_type)
            sample_sizes.update(samples_data.keys())
        
        sample_sizes = sorted(sample_sizes)
        model_types = sorted(model_types)
        
        if sample_sizes and model_types:
            table = doc.add_table(rows=len(model_types)+1, cols=len(sample_sizes)+1)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Model Type'
            for i, samples in enumerate(sample_sizes, 1):
                header_cells[i].text = f'{samples} samples'
            
            # Data rows
            for row_idx, model_type in enumerate(model_types, 1):
                row_cells = table.rows[row_idx].cells
                row_cells[0].text = model_type
                
                for col_idx, samples in enumerate(sample_sizes, 1):
                    error = model_data[model_type].get(samples, float('inf'))
                    if error == float('inf'):
                        row_cells[col_idx].text = 'FAIL'
                    else:
                        # Convert to percentage error (assuming relative error)
                        error_percent = error * 100
                        row_cells[col_idx].text = f'{error_percent:.2f}%'
            
            # Create graph
            create_error_graph(func_name, model_data, sample_sizes, model_types)
            graph_path = f'{func_name}_error_graph.png'
            doc.add_picture(graph_path, width=Inches(6))
            doc.add_paragraph(f'Figure: Error trends for {func_name} function')
            
        doc.add_paragraph()  # Add spacing between functions
    
    # Save document
    doc.save('AI_Models_Test_Report.docx')
    print(f"\nWord report saved as 'AI_Models_Test_Report.docx'")

def create_error_graph(func_name, model_data, sample_sizes, model_types):
    """Create double logarithmic graph of errors"""
    plt.figure(figsize=(10, 6))
    
    for model_type in model_types:
        errors = []
        samples_list = []
        
        for samples in sample_sizes:
            error = model_data[model_type].get(samples, float('inf'))
            if error < float('inf'):
                errors.append(error)
                samples_list.append(samples)
        
        if errors:
            plt.loglog(samples_list, errors, 'o-', label=model_type, linewidth=2, markersize=8)
    
    plt.xlabel('Number of Samples', fontsize=12)
    plt.ylabel('Mean Absolute Error', fontsize=12)
    plt.title(f'Error vs Sample Size for {func_name}', fontsize=14)
    plt.legend()
    plt.grid(True, which="both", ls="-", alpha=0.2)
    plt.tight_layout()
    
    # Save graph
    plt.savefig(f'{func_name}_error_graph.png', dpi=300, bbox_inches='tight')
    plt.close()

import multiprocessing
from concurrent.futures import ProcessPoolExecutor, as_completed
import os
import time
from collections import defaultdict

def test_model_wrapper(model_file):
    if not os.path.exists(model_file):
        print(f"WARNING: Model file not found: {model_file}")
        return None, None, None, float('inf')
    
    from test_report import parse_model_filename, test_single_model
    
    func_name, model_type, samples = parse_model_filename(model_file)
    error = test_single_model(model_file)
    return func_name, model_type, samples, error

def main():
    print("AI Models Test Report - Real Error Evaluation")
    print("=" * 60)
    
    if not models_info:
        print("No models to test!")
        return
    
    # Organize results by function and model type
    results = defaultdict(lambda: defaultdict(dict))
    
    print(f"Testing {len(models_info)} models on 1000 random points each...")
    print(f"Using {multiprocessing.cpu_count()} CPU cores")
    print()
    
    start_time = time.time()
    
    max_workers = max(1, multiprocessing.cpu_count() - 1)
    print(f"Using {max_workers} workers")
    
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        future_to_model = {
            executor.submit(test_model_wrapper, model_file): model_file 
            for model_file in models_info
        }
        
        completed = 0
        for future in as_completed(future_to_model):
            model_file = future_to_model[future]
            try:
                func_name, model_type, samples, error = future.result()
                if func_name is not None:
                    results[func_name][model_type][samples] = error
                
                completed += 1
                print(f"Progress: {completed}/{len(models_info)} - {func_name} {model_type} {samples}samples: {error:.6f}")
                
            except Exception as e:
                print(f"ERROR processing {model_file}: {e}")
                completed += 1
    
    end_time = time.time()
    total_time = end_time - start_time
    print(f"\nTotal testing time: {total_time:.2f} seconds")
    
    # Print console summary
    print("\n" + "=" * 60)
    print("SUMMARY REPORT")
    print("=" * 60)
    
    for func_name, model_data in sorted(results.items()):
        if func_name not in function_info:
            continue
            
        func_info = function_info[func_name]
        input_dims = func_info['input_dims']
        output_dims = func_info['output_dims']
        
        print(f"\nFunction: {func_name} ({input_dims}D -> {output_dims}D)")
        print("-" * 50)
        
        sample_sizes = set()
        model_types = set()
        
        for model_type, samples_data in model_data.items():
            model_types.add(model_type)
            sample_sizes.update(samples_data.keys())
        
        sample_sizes = sorted(sample_sizes)
        model_types = sorted(model_types)
        
        # Print header
        header = "Model Type".ljust(12)
        for samples in sample_sizes:
            header += f"{samples}samples".ljust(12)
        print(header)
        print("-" * 50)
        
        # Print data for each model type
        for model_type in model_types:
            row = model_type.ljust(12)
            for samples in sample_sizes:
                error = model_data[model_type].get(samples, float('inf'))
                if error == float('inf'):
                    row += "FAIL".ljust(12)
                else:
                    error_percent = error * 100
                    row += f"{error_percent:.2f}%".ljust(12)
            print(row)
    
    # Create Word report
    print("\nGenerating Word report...")
    create_word_report(results, total_time)
    
    print("\n" + "=" * 60)
    print("TESTING COMPLETED")
    print("=" * 60)

if __name__ == '__main__':
    main()